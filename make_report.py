# make_report.py
from docx import Document
from pathlib import Path
import json

root = Path(".").resolve()
doc = Document()
doc.add_heading("Assignment 3 – SeqTrack Setup, Training, and Checkpoint Management", level=1)
doc.add_paragraph("Team number: 7")
doc.add_paragraph("Team leader Hugging Face repo: https://huggingface.co/saifmamdouh11/seqtrack_assignment3")
doc.add_paragraph("GitHub repo: https://github.com/saif-mamdouh/assignment-3-")
doc.add_heading("Selected classes and dataset sizes", level=2)

try:
    subset = json.load(open(root/"lasot_subset"/"subset_info.json"))
    for cls,info in subset.items():
        doc.add_paragraph(f"{cls}: train_size={info['train_size']}, test_size={info['test_size']}")
except Exception as e:
    doc.add_paragraph("Could not read lasot_subset/subset_info.json: " + str(e))

doc.add_heading("Environment & Installed packages", level=2)
doc.add_paragraph("See installed_packages.txt attached.")

doc.add_heading("Code modifications (high-level)", level=2)
doc.add_paragraph("Files modified/added and locations (update line numbers if code changed):")
doc.add_paragraph(" - checkpoint_utils.py : save/load checkpoint including optimizer/scheduler/RNG states")
doc.add_paragraph(" - train.py : main training loop, seed set at epoch start, per-epoch checkpoint save & HF upload, --resume_from support")
doc.add_paragraph(" - utils/logger.py : prints time stats every 50 samples and writes training log")
doc.add_paragraph(" - seqtrack_mods/model_wrapper.py : wrapper to import official SeqTrack or fallback toy model")
doc.add_paragraph(" - compare_metrics.py & plot_metrics.py & make_report.py : compare, plot, and generate docx")

doc.add_heading("Training results", level=2)
doc.add_paragraph("Include loss_phase1.png, iou_phase1.png and the phase2 counterparts here.")
doc.add_heading("Reproducibility notes", level=2)
doc.add_paragraph("Seed: team number 7. Deterministic torch backend flags set. Use same environment (installed_packages.txt) to reproduce.")

doc.save(root/"assignment_3.docx")
print("Saved assignment_3.docx")
